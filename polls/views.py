from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from .forms import UploadFileForm
import openai
import json
from django.core.files.storage import default_storage
import docx
import os
from pathlib import Path
from rest_framework.decorators import api_view
from rest_framework.response import Response

openai.api_key = "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"


# Create your views here.
def index(request):
	return HttpResponse("Hello world, youre at polls index")


def get_completion(prompt):
	print("Prompt:")
	print(prompt)

	query = openai.Completion.create(
		 engine="text-davinci-003",
		 prompt=prompt,
		 max_tokens=1024,
		 n=1,
		 stop=None,
		 temperature=0.5,
	)

	response = query.choices[0].text
	return response

def query_view(request):
	if request.method == "POST":
		print(request)

		f = request.FILES['sentFile']
		# print(f)
		response = {}
		file_name = "doc.docx"
		file_name_2 = default_storage.save(file_name, f)
		url = default_storage.url(file_name_2)
		print(url)
		# print(os.path.dirname(__file__))

		parent = Path(__file__).parent.parent
		print(parent)

		path = parent / "media" / file_name_2
		print(path)


		# print(file_url)
		# return render(request, 'index.html', response)
		doc = docx.Document(path)
		fullText = []
		for para in doc.paragraphs:
			fullText.append(para.text)

		prompt = '\n'.join(fullText) + " Based on the above notes, generate 3 multiple choice question with 4 answer choices. Please give the correct answers at the end of your response in a separate section. For example: Q1: blah blah Q2: blah blah Correct Answers: blah blah"

		response = get_completion(prompt)
		spl = response.split("Correct Answers:")
		question = spl[0].strip()
		answer = spl[1].strip()
		responseJson = {"question": question, "answer": answer}
		print(json.dumps(responseJson))
		responseJson['name'] = str(response)
		# print(responseJson)
		return render(request, 'index.html', responseJson)
	return render(request, 'index.html')

@api_view(['POST', 'GET'])
def process_text_view(request):
	if (request.method == "GET"):
		return render(request, 'index.html')

	# if (request.method == 'OPTIONS'):
	# 	http_method_names = ['GET', 'POST']
	# 	print(request)
	# 	print(request.path)
	# 	print(request.path_info)
	# 	print(request.FILES)
	# 	print(request.META)
	# 	response = HttpResponse()
	# 	response['allow'] = http_method_names
	# 	return response

	print(request)
	print(request.POST)
	print(request.FILES)
	
	f = request.FILES['sentFile']
	# print(f)
	response = {}
	file_name = "doc.docx"
	file_name_2 = default_storage.save(file_name, f)
	url = default_storage.url(file_name_2)
	print(url)
	# print(os.path.dirname(__file__))

	parent = Path(__file__).parent.parent
	print(parent)

	path = parent / "media" / file_name_2
	print(path)


	# print(file_url)
	# return render(request, 'index.html', response)
	doc = docx.Document(path)
	fullText = []
	for para in doc.paragraphs:
		fullText.append(para.text)

	prompt = '\n'.join(fullText) + " Based on the above notes, generate 3 multiple choice question with 4 answer choices. Please give the correct answers at the end of your response in a separate section. For example: Q1: blah blah Q2: blah blah Correct Answers: blah blah"

	responseStr = get_completion(prompt)
	spl = responseStr.split("Correct Answers:")
	question = spl[0].strip()
	answer = spl[1].strip()
	responseJson = {"question": question, "answer": answer}
	


	jsonObj = json.dumps(responseJson)

	jsonPath = parent / "note-ai-frontend" / "src" / "lib" / "data" / "testPrompt.json"

	with open(jsonPath, "w") as outfile:
		outfile.write(jsonObj)


	# responseJson['name'] = str(response)
	# print(responseJson)
	response = Response(responseJson)
	response["Access-Control-Allow-Origin"] = "*"
	response["Access-Control-Allow-Methods"] = "GET, OPTIONS"
	response["Access-Control-Max-Age"] = "1000"
	response["Access-Control-Allow-Headers"] = "X-Requested-With, Content-Type"
	return response